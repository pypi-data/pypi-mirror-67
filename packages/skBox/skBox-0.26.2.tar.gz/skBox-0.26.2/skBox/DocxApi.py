from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pandas as pd


class Docxer(object):
    """
    python-docx Word文档API再封装,用于非常便捷得使用python代码编写word文档
    """
    def __init__(self):
        # 生成doc文件
        self.doc = Document()
        # 更改全局字体
        self.doc.styles['Normal'].font.name = u'宋体'
        self.doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        # 生成paragraph对象
        self.p = None

    def add_paragraph(self, input_str, alignment=None, ctn=False):
        """
        增加段落
        ctn如果为True就代表续写paragraph
        :param input_str: 输入内容
        :param alignment: 是否首行缩进
        :param ctn: False->重开段落， True->续写段落
        :return:
        """
        if not ctn:
            self.p = self.doc.add_paragraph(input_str)
            if alignment == '居中':
                self.p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                pass
            self.p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            self.p.paragraph_format.space_after = 1
        else:
            self.p.add_run(input_str)

    def add_heading(self, input_str, level=0):
        self.doc.add_heading(input_str, level=level)

    def add_table(self, df, fontsize=None, index=False):
        """
        输入pands.Datafram，输出word表格
        :param df: data_frame
        :param fontsize: 字体大小
        :return:
        """
        if index:
            df = df.reset_index()
        row_nums, columns_num = df.shape[0] + 1, df.shape[1]
        df = df.fillna('').astype(str)
        t = self.doc.add_table(row_nums, columns_num)
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        doc_columns = list(df.columns)
        print(doc_columns)

        # add the header rows.
        for j in range(0, columns_num):
            t.cell(0, j).text = doc_columns[j]

        # add the rest of the data frame
        for i in range(1, row_nums):
            for j in range(0, columns_num):
                t.cell(i, j).text = str(df.iloc[i-1, j])

        # 设置字体大小
        if fontsize is not None:
            for row in t.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(fontsize)

    def add_picture(self, dir, width=6, height=None):
        """
        在当前位置添加图片
        :param dir: 图片地址
        :param width: 宽度
        :param height: 高度
        :return:
        """
        self.doc.add_picture(dir, width=Inches(width), height=Inches(height))

    def save(self, path):
        self.doc.save(path+".docx")


if __name__ == '__main__':
    doc = Docxer()
    doc.add_heading("我是谁")
    doc.add_paragraph("今天我去哪了", alignment='居中')
    doc.add_paragraph("我明天去哪")
    df = pd.DataFrame({"aa": range(10), "bb":range(10, 20)}, index=range(-10, 0))
    df.index.name = '我是索引'
    doc.add_table(df, index=True)
    doc.save("测试")

