import re

import docx, xlwings

class TemplateDocument(object):
    def __init__(self, document, keywords=None):
        self.document = docx.Document(document)
        self.keywords = keywords or self.detect_keywords()
        self.run_maps = {}

    def parse_document(self):
        run_maps = []

        self._parse_paragraphs(self.document.paragraphs, run_maps)
        self._parse_tables(self.document.tables, run_maps)
        self._parse_sections(self.document.sections, run_maps)

        self.run_maps = self._join_run_maps(run_maps)
        return self.run_maps

    def _parse_paragraphs(self, paragraphs, rm):
        for p in paragraphs:
            m = self._parse_paragraph(p)
            if m:
                rm.append(m)

    def _parse_tables(self, tables, rm):
        for table in tables:
            for row in table.rows:
                for c in row.cells:
                    for p in c.paragraphs:
                        m = self._parse_paragraph(p)
                        if m:
                            rm.append(m)
                    if c.tables:
                        self._parse_tables(c.tables, rm)

    def _parse_sections(self, sections, rm):
        for s in sections:
            self._parse_paragraphs(s.header.paragraphs, rm)
            self._parse_paragraphs(s.footer.paragraphs, rm)
            self._parse_paragraphs(s.even_page_header.paragraphs, rm)
            self._parse_paragraphs(s.even_page_footer.paragraphs, rm)
            self._parse_paragraphs(s.first_page_header.paragraphs, rm)
            self._parse_paragraphs(s.first_page_footer.paragraphs, rm)

            self._parse_tables(s.header.tables, rm)
            self._parse_tables(s.footer.tables, rm)
            self._parse_tables(s.even_page_header.tables, rm)
            self._parse_tables(s.even_page_footer.tables, rm)
            self._parse_tables(s.first_page_header.tables, rm)
            self._parse_tables(s.first_page_footer.tables, rm)

    def _parse_paragraph(self, paragraph):
        '''
        doc:
            abcdefgh
           (01234567)
            ^  ^  ^
        run:|  |  |
            0  1  2
        run_dict:
            { 0: run0, 3: run1, 6: run2 }
        result:
            { 
                'cdefg': [(run0, 'ab')],
            }
            run1.text == ''
            run2.text == 'h'
        '''

        # split rare paragraph into run_dict
        run_count, run_dict = 0, {}
        content = []
        result = {}

        for r in paragraph.runs:

            t = r.text

            # join texts
            content.append(t)
            # sort dict
            run_dict[run_count] = r
            run_count += len(t)

        content = ''.join(content)

        for keyword in self.keywords:

            # find keyword matches and get start points
            c, keys = content, run_dict.keys()
            matches, start_point = {}, 0

            while True:
                r = re.search('{{ %s }}' % keyword, c)
                if r is None:
                    break

                # get match run
                match = []
                start = start_point + r.start()
                while start not in keys:
                    start -= 1
                match.append(run_dict[start])
                for i in range(r.start() + 1, r.end()):
                    i += start_point
                    if i in keys:
                        match.append(run_dict[i])

                # format match run
                if 1 < len(match):
                    match[0].text = re.sub(
                        '{?{[^{]*$',
                        '{{ %s }}' % keyword,
                        match[0].text)
                    for m in match[1:-1]:
                        m.text = ''
                    match[-1].text = re.sub('^.*?}}?', '', match[-1].text)

                matches[match[0]] = match[0].text
                # match[0].text can be 'blabla{{ keyword }}' or 'blabla{{ keyword }}blabla'

                c = c[r.end():]
                start_point += r.end()

            if matches:
                result[keyword] = [i for i in matches.items()]

        return result

    def _join_run_maps(self, run_maps):
        r = {}
        for m in run_maps:
            for k, v in m.items():
                if k not in r:
                    r[k] = []
                r[k].extend(v)
        return r

    def replace_and_save(self, value_map, filename):
        for k, v in value_map.items():
            for run, _ in self.run_maps.get(k, []):
                run.text = re.sub('{{ %s }}' % k, v, run.text)
        
        self.document.save(filename)

        for runs in self.run_maps.values():
            for run, text in runs:
                run.text = text

    def clear_user_info(self):
        return


def load_excel(filename):
    app = xlwings.App(visible=False)
    book = app.books.open(filename)
    sheet = book.sheets[0]
    line, value_map = 1, {}

    try:
        while True:
            key = sheet.range('A%s' % line).value
            if not key:
                break
            value = sheet.range('B%s' % line).value
            value_map[key] = value
            line += 1
    except:
        pass
    finally:
        book.close()

    return value_map

if __name__ == '__main__':
    vm = {
        'topic': 'my-topic',
        'content': 'my-content',
    }
    template = TemplateDocument('template.docx', vm.keys())
    template.parse_document()
    template.replace_and_save(vm, 'outcome.docx')

    vm = {
        'topic': 'my-another-topic',
        'content': 'my-another-content',
    }
    template.replace_and_save(vm, 'another-outcome.docx')
